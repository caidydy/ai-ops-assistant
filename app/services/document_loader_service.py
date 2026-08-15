"""多格式文档加载服务模块

负责把不同格式的知识库文件解析为纯文本，供文档分割器统一处理。

支持格式：
- .md / .txt    直接读取 UTF-8 文本
- .pdf          PyMuPDF 逐页提取文本（中文排版支持好）
- .docx         python-docx 提取段落文本
- .xlsx         openpyxl 按工作表提取单元格值（保留表头/行号语义）

对外只暴露 load(file_path) -> str 单一接口，调用方无需关心文件格式。
"""

from pathlib import Path
from typing import Set

from loguru import logger

# 支持的文件扩展名（小写，不含点），与 app/api/file.py 的 ALLOWED_EXTENSIONS 保持一致
SUPPORTED_EXTENSIONS: Set[str] = {"md", "txt", "pdf", "docx", "xlsx"}


class DocumentLoaderService:
    """文档加载服务 - 按扩展名分发到对应解析器"""

    def load(self, file_path: str) -> str:
        """
        加载文件并返回纯文本内容

        Args:
            file_path: 文件路径

        Returns:
            str: 提取到的文本内容

        Raises:
            ValueError: 文件不存在或不支持的格式
        """
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            raise ValueError(f"文件不存在: {file_path}")

        extension = path.suffix.lower().lstrip(".")
        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"不支持的文件格式: .{extension}，仅支持: "
                f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        logger.debug(f"开始加载文档: {file_path} (格式: .{extension})")

        loaders = {
            "md": self._load_text,
            "txt": self._load_text,
            "pdf": self._load_pdf,
            "docx": self._load_docx,
            "xlsx": self._load_xlsx,
        }
        content = loaders[extension](path)

        logger.debug(f"文档加载完成: {file_path}, 文本长度: {len(content)} 字符")
        return content

    # ---------------------------------------------------------------
    # 各格式解析器
    # ---------------------------------------------------------------

    @staticmethod
    def _load_text(path: Path) -> str:
        """读取 UTF-8 纯文本（md / txt）"""
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            # 兜底：部分 txt 可能是 GBK/GB18030 编码（常见于中文 Windows 产物）
            logger.warning(f"UTF-8 解码失败，尝试 GB18030: {path}")
            return path.read_text(encoding="gb18030")

    @staticmethod
    def _load_pdf(path: Path) -> str:
        """使用 PyMuPDF 逐页提取 PDF 文本"""
        import fitz  # pymupdf

        parts: list[str] = []
        with fitz.open(path) as doc:
            for page in doc:
                text = page.get_text()
                if text and text.strip():
                    parts.append(text.strip())
        if not parts:
            logger.warning(
                f"PDF 未提取到文本（可能是扫描件/纯图片，需 OCR 支持）: {path}"
            )
        # 页与页之间用空行分隔，避免正文粘连
        return "\n\n".join(parts)

    @staticmethod
    def _load_docx(path: Path) -> str:
        """使用 python-docx 提取 Word 段落文本"""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        # 表格内容同样有价值（如参数对照表）
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts)

    @staticmethod
    def _load_xlsx(path: Path) -> str:
        """使用 openpyxl 提取 Excel 单元格值（每个工作表一段表格化文本）"""
        from openpyxl import load_workbook

        wb = load_workbook(str(path), read_only=True, data_only=True)
        sections: list[str] = []
        for sheet in wb.worksheets:
            lines: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                # 跳过整行为空的行
                values = [str(v) if v is not None else "" for v in row]
                if any(v.strip() for v in values):
                    lines.append("\t".join(values))
            if lines:
                # 带上工作表名作为标题，保留表格语义
                sections.append(f"[工作表: {sheet.title}]\n" + "\n".join(lines))
        wb.close()
        return "\n\n".join(sections)


# 全局单例
document_loader_service = DocumentLoaderService()